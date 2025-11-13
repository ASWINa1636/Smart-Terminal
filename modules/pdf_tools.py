from pathlib import Path
from PyPDF2 import PdfReader, PdfWriter
from docx import Document
from docxcompose.composer import Composer
from docx2pdf import convert
from PIL import Image
from rich.console import Console
from rich.prompt import Prompt
import subprocess
import platform

console = Console()

def convert_doc_to_docx(file_path: Path):
    console.print(f"Converting old Word file: {file_path.name}")
    if not file_path.exists() or file_path.suffix.lower() != ".doc":
        console.print("[red]Not a .doc file[/red]")
        return None

    out_path = file_path.with_suffix(".docx")

    try:
        if platform.system() == "Windows":
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            doc = word.Documents.Open(str(file_path))
            doc.SaveAs(str(out_path), FileFormat=16)  
            doc.Close()
            word.Quit()
        else:
            subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "docx", str(file_path)],
                check=True
            )
        console.print(f"✅ Converted to [green]{out_path.name}[/green]")
        return out_path
    except Exception as e:
        console.print(f"[red]Failed to convert {file_path.name}: {e}[/red]")
        return None


def select_file_from_folder(path: Path, file_types: list[str], description="file"):
    path = path.resolve()
    if path.is_file():
        return [path]
    elif path.is_dir():
        files = [f for f in path.iterdir() if f.suffix.lower() in file_types]
        if not files:
            console.print(f"[yellow]No {description}s found in this folder.[/yellow]")
            return []
        console.print(f"\n[bold cyan]Select {description}(s) to process:[/bold cyan]")
        for i, f in enumerate(files, start=1):
            console.print(f"{i}. {f.name}")
        console.print("\n[dim]Example: 1,3,5 or 2-6 or all[/dim]")
        selection = Prompt.ask(f"Enter number(s) (1-{len(files)})", default="1").strip().lower()

        selected = []
        try:
            if selection == "all":
                return files
            parts = [s.strip() for s in selection.split(",")]
            for p in parts:
                if "-" in p:
                    start, end = map(int, p.split("-"))
                    selected.extend(files[start - 1:end])
                else:
                    idx = int(p)
                    selected.append(files[idx - 1])
        except Exception:
            console.print("[red]Invalid input, selecting first file instead.[/red]")
            selected = [files[0]]
        return selected
    else:
        console.print(f"[red]Invalid path[/red]")
        return []


def merge_pdfs():
    console.print("\n[bold cyan]Merge PDFs[/bold cyan]")
    folder_input = Prompt.ask("Enter folder containing PDFs", default=".")
    folder = Path(folder_input).expanduser().resolve()

    if not folder.exists() or not folder.is_dir():
        console.print(f"[red]Folder not found:[/red] {folder}")
        console.print("[yellow]Tip: Check the folder name and try again (e.g., /home/aswin/Documents)[/yellow]\n")
        return

    pdfs = [f for f in folder.iterdir() if f.suffix.lower() == ".pdf"]
    if not pdfs:
        console.print("[yellow]No PDF files found in this folder.[/yellow]\n")
        return

    console.print("\n[bold cyan]Select PDF(s) to merge:[/bold cyan]")
    for i, pdf in enumerate(pdfs, start=1):
        console.print(f"{i}. {pdf.name}")
    console.print("\n[dim]Example: 1,3,5 or 2-6 or all[/dim]")

    selection = Prompt.ask(f"Enter number(s) (1-{len(pdfs)})", default="all").strip().lower()

    selected = []
    try:
        if selection == "all":
            selected = pdfs
        else:
            parts = [s.strip() for s in selection.split(",")]
            for p in parts:
                if "-" in p:
                    start, end = map(int, p.split("-"))
                    selected.extend(pdfs[start - 1:end])
                else:
                    idx = int(p)
                    selected.append(pdfs[idx - 1])
    except Exception:
        console.print("[red]Invalid input, selecting all PDFs instead.[/red]")
        selected = pdfs

    if not selected:
        console.print("[yellow]No PDFs selected.[/yellow]")
        return

    output_name = Prompt.ask("Enter output filename", default="merged.pdf")
    out = folder / output_name

    writer = PdfWriter()
    for f in selected:
        try:
            reader = PdfReader(f)
            for page in reader.pages:
                writer.add_page(page)
            console.print(f"✅ Added {f.name}")
        except Exception as e:
            console.print(f"[red]Failed to add {f.name}: {e}[/red]")

    # Save merged file
    with open(out, "wb") as f:
        writer.write(f)
    console.print(f"\nMerged PDF saved: [green]{out}[/green]\n")



def split_pdf():
    console.print("\n[bold cyan]Split PDF[/bold cyan]")
    selected_files = select_file_from_folder(
        Path(Prompt.ask("Enter PDF file or folder", default=".")),
        [".pdf"],
        "PDF"
    )
    if not selected_files:
        return

    for file in selected_files:
        try:
            reader = PdfReader(file)
            total = len(reader.pages)
            console.print(f"\n{file.name} - Total Pages: {total}")
            console.print("[dim]Example: 1-3,6-8 or 2,4,9[/dim]")

            ranges = Prompt.ask("Enter page ranges to extract", default=f"1-{total}").strip()

            # Parse multi-ranges like 1-3,6-8
            page_numbers = set()
            try:
                for part in ranges.split(","):
                    part = part.strip()
                    if "-" in part:
                        start, end = map(int, part.split("-"))
                        page_numbers.update(range(start, end + 1))
                    else:
                        page_numbers.add(int(part))
            except Exception:
                console.print("[red]Invalid range format. Example: 1-3,5,7-9[/red]")
                continue

            # Create a new PDF for selected pages
            writer = PdfWriter()
            for i in sorted(page_numbers):
                if 1 <= i <= total:
                    writer.add_page(reader.pages[i - 1])

            # Output file name
            out = file.parent / f"{file.stem}_split_selected.pdf"
            with open(out, "wb") as f:
                writer.write(f)

            console.print(f"✅ Split file created: [green]{out.name}[/green]\n")

        except Exception as e:
            console.print(f"[red]Failed to split {file.name}: {e}[/red]")


def word_to_pdf():
    console.print("\n[bold cyan]Word → PDF Converter[/bold cyan]")
    selected = select_file_from_folder(
        Path(Prompt.ask("Enter Word file or folder", default=".")),
        [".doc", ".docx"],
        "Word file"
    )
    if not selected:
        return

    for file in selected:
        # Auto-convert .doc → .docx if needed
        if file.suffix.lower() == ".doc":
            new_path = convert_doc_to_docx(file)
            if new_path:
                file = new_path
            else:
                console.print(f"[red]Skipping {file.name} (conversion failed)[/red]")
                continue

        # Output PDF in the same folder as the Word file
        output_path = file.parent / f"{file.stem}.pdf"

        # Platform-specific conversion
        if platform.system() == "Windows":
            try:
                from docx2pdf import convert
                convert(str(file), str(output_path))
                console.print(f"Saved [green]{output_path.name}[/green] in {file.parent}")
            except Exception as e:
                console.print(f"[red]Failed to convert {file.name}: {e}[/red]")
        else:
            # Linux / macOS → use LibreOffice headless mode
            try:
                subprocess.run(
                    [
                        "libreoffice",
                        "--headless",
                        "--convert-to", "pdf",
                        "--outdir", str(file.parent),
                        str(file)
                    ],
                    check=True
                )
                console.print(f"Saved [green]{output_path.name}[/green] in {file.parent}")
            except Exception as e:
                console.print(f"[red]LibreOffice conversion failed for {file.name}: {e}[/red]")


def image_to_pdf():
    console.print("\n[bold cyan]Image → PDF Converter[/bold cyan]")
    path = Path(Prompt.ask("Enter image file or folder", default=".")).resolve()

    # If it's a single file
    if path.is_file():
        if path.suffix.lower() not in [".jpg", ".jpeg", ".png"]:
            console.print("[red]Not an image file[/red]")
            return
        image = Image.open(path).convert("RGB")
        out = path.with_suffix(".pdf")
        image.save(out)
        console.print(f"Converted [green]{path.name}[/green] → [green]{out.name}[/green]\n")
        return

    # If it's a folder
    if path.is_dir():
        imgs = [f for f in path.iterdir() if f.suffix.lower() in [".jpg", ".jpeg", ".png"]]
        if not imgs:
            console.print("[yellow]No images found in this folder.[/yellow]")
            return

        console.print("\n[bold cyan]Select images to include in PDF:[/bold cyan]")
        for i, img in enumerate(imgs, start=1):
            console.print(f"{i}. {img.name}")
        console.print("\n[dim]Example: 1,3,5 or 2-6 or all[/dim]")

        selection = Prompt.ask(f"Enter number(s) (1-{len(imgs)})", default="all").strip().lower()

        selected = []
        try:
            if selection == "all":
                selected = imgs
            else:
                parts = [s.strip() for s in selection.split(",")]
                for p in parts:
                    if "-" in p:
                        start, end = map(int, p.split("-"))
                        selected.extend(imgs[start - 1:end])
                    else:
                        idx = int(p)
                        selected.append(imgs[idx - 1])
        except Exception:
            console.print("[red]Invalid input, selecting all images instead.[/red]")
            selected = imgs

        if not selected:
            console.print("[yellow]No images selected.[/yellow]")
            return

        # Open and merge selected images
        images = [Image.open(img).convert("RGB") for img in selected]
        output_name = Prompt.ask("Enter output PDF name", default="images_merged.pdf")
        out = path / output_name
        images[0].save(out, save_all=True, append_images=images[1:])

        console.print(f"PDF created: [green]{out}[/green]\n")


def merge_word():
    console.print("\n[bold cyan]Merge Word Files[/bold cyan]")
    folder = Path(Prompt.ask("Enter folder containing Word files", default=".")).resolve()
    docs = [f for f in folder.iterdir() if f.suffix.lower() in [".doc", ".docx"]]
    if not docs:
        console.print("[yellow]No Word files found.[/yellow]")
        return

    # Convert any .doc files to .docx first
    converted_docs = []
    for f in docs:
        if f.suffix.lower() == ".doc":
            new = convert_doc_to_docx(f)
            if new:
                converted_docs.append(new)
        else:
            converted_docs.append(f)

    if not converted_docs:
        console.print("[red]No valid Word files to merge.[/red]")
        return

    merged = Document(converted_docs[0])
    composer = Composer(merged)
    for f in converted_docs[1:]:
        try:
            composer.append(Document(f))
            console.print(f"✅ Added {f.name}")
        except Exception as e:
            console.print(f"[red]Failed to merge {f.name}: {e}[/red]")

    out = folder / "merged.docx"
    composer.save(out)
    console.print(f"Merged Word file saved: [green]{out}[/green]\n")



def split_word():
    console.print("\n[bold cyan]Split Word File[/bold cyan]")
    selected_files = select_file_from_folder(
        Path(Prompt.ask("Enter Word file or folder", default=".")),
        [".docx"],
        "Word file"
    )
    if not selected_files:
        return

    for file in selected_files:
        console.print(f"\nProcessing: [green]{file.name}[/green]")
        try:
            doc = Document(file)
            if not hasattr(doc, "paragraphs"):
                raise ValueError("File opened but not a valid Word document.")
        except Exception as e:
            console.print(f"[red]Failed to open {file.name}: {e}[/red]")
            continue

        paras = doc.paragraphs
        if not paras:
            console.print(f"[yellow] {file.name} has no paragraphs to split.[/yellow]")
            continue

        chunk = int(Prompt.ask("Split after how many paragraphs?", default="5"))
        out_dir = file.parent / f"{file.stem}_split"
        out_dir.mkdir(exist_ok=True)

        parts = [paras[i:i + chunk] for i in range(0, len(paras), chunk)]
        for i, part in enumerate(parts, start=1):
            new_doc = Document()
            for p in part:
                new_doc.add_paragraph(p.text)
            out_path = out_dir / f"{file.stem}_part{i}.docx"
            new_doc.save(out_path)
            console.print(f"✅ Created {out_path.name}")

        console.print(f"🎉 Split files saved in [green]{out_dir}[/green]\n")


def protect_pdf():
    console.print("\n[bold cyan]Protect PDF File[/bold cyan]")
    selected_files = select_file_from_folder(Path(Prompt.ask("Enter PDF file or folder", default=".")), [".pdf"], "PDF")
    if not selected_files:
        return

    password = Prompt.ask("Enter password to protect")

    for file in selected_files:
        reader = PdfReader(file)
        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)
        writer.encrypt(password)
        out = file.parent / f"{file.stem}_protected.pdf"
        with open(out, "wb") as f:
            writer.write(f)
        console.print(f"✅ Protected: [green]{out.name}[/green]")

    console.print(f"\nProtected {len(selected_files)} file(s) successfully!\n")


def unlock_pdf():
    console.print("\n[bold cyan]Unlock PDF File (remove password)[/bold cyan]")
    selected_files = select_file_from_folder(
        Path(Prompt.ask("Enter PDF file or folder", default=".")),
        [".pdf"],
        "PDF"
    )
    if not selected_files:
        return

    password = Prompt.ask("Enter password used for the PDF(s)", password=True)
    unlocked_count = 0

    for file in selected_files:
        try:
            reader = PdfReader(file)
        except Exception as e:
            console.print(f"[red]Could not open {file.name}: {e}[/red]")
            continue

        if not getattr(reader, "is_encrypted", False):
            console.print(f"[yellow]{file.name} is not encrypted — skipping.[/yellow]")
            continue

        try:
            reader.decrypt(password)
            try:
                _ = reader.pages[0]
                can_read = True
            except Exception:
                can_read = False

            if not can_read:
                console.print(f"[red]Wrong password for {file.name} — skipping.[/red]")
                continue

            writer = PdfWriter()
            for page in reader.pages:
                writer.add_page(page)

            out_path = file.parent / f"{file.stem}_unlocked.pdf"
            with open(out_path, "wb") as out_f:
                writer.write(out_f)
            console.print(f"✅ Unlocked: [green]{out_path.name}[/green]")
            unlocked_count += 1

        except Exception as e:
            console.print(f"[red]Failed to unlock {file.name}: {e}[/red]")

    console.print(f"\nDone — unlocked {unlocked_count} file(s).\n")


def pdf_menu():
    while True:
        console.print("""
[bold cyan]PDF & Word Automation Tools[/bold cyan]
1. Merge PDFs
2. Split PDF
3. Word → PDF
4. Image → PDF
5. Merge Word Files
6. Split Word File
7. Protect PDF
8. Unlock PDF
9. Back to Main Menu
""")
        choice = Prompt.ask("[bold green]Choose option[/bold green]", choices=[str(i) for i in range(1, 10)])
        if choice == "1":
            merge_pdfs()
        elif choice == "2":
            split_pdf()
        elif choice == "3":
            word_to_pdf()
        elif choice == "4":
            image_to_pdf()
        elif choice == "5":
            merge_word()
        elif choice == "6":
            split_word()
        elif choice == "7":
            protect_pdf()
        elif choice == "8":
            unlock_pdf()
        elif choice == "9":
            return
        Prompt.ask("\nPress Enter to continue...")
